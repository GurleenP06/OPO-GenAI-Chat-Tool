from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
from pydantic import BaseModel
from typing import List, Optional, Dict
import config
import rag_model
import uuid
import json
import os
from datetime import datetime
import pandas as pd
from pathlib import Path
import re

# Document generation imports
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
import io

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Data storage paths
DATA_DIR = Path("./chat_data")
DATA_DIR.mkdir(exist_ok=True)
PROJECTS_FILE = DATA_DIR / "projects.json"
CHATS_FILE = DATA_DIR / "chats.json"
RATINGS_FILE = DATA_DIR / "ratings.json"

# Initialize data files
for file_path in [PROJECTS_FILE, CHATS_FILE, RATINGS_FILE]:
    if not file_path.exists():
        with open(file_path, "w") as f:
            json.dump({} if file_path != RATINGS_FILE else [], f)

# In-memory storage
chat_sessions = {}
projects = {}
chat_metadata = {}

# Load existing data
def load_data():
    global projects, chat_metadata
    try:
        with open(PROJECTS_FILE, "r") as f:
            projects = json.load(f)
        with open(CHATS_FILE, "r") as f:
            chat_metadata = json.load(f)
    except:
        pass

load_data()

# Models
class Project(BaseModel):
    name: str
    description: Optional[str] = ""

class ChatMetadata(BaseModel):
    session_id: str
    name: str
    project_id: Optional[str] = None
    is_favorite: bool = False
    created_at: str
    updated_at: str

class RenameRequest(BaseModel):
    session_id: str
    new_name: str

class MoveToProjectRequest(BaseModel):
    session_id: str
    project_id: Optional[str]

class ToggleFavoriteRequest(BaseModel):
    session_id: str

class ExportRequest(BaseModel):
    session_id: str
    message_index: int  # Index of the message to export
    format: str  # 'docx' or 'pdf'

class QueryRequest(BaseModel):
    session_id: str
    query: str

class ChatHistoryRequest(BaseModel):
    session_id: str

class RatingRequest(BaseModel):
    question: str
    response: str
    feedback_type: str  # 'positive' or 'negative'
    selected_reason: Optional[str] = None
    custom_feedback: Optional[str] = None

class DocumentViewRequest(BaseModel):
    filename: str
    highlights: List[Dict]  # List of passages to highlight

# Project endpoints
@app.post("/create_project/")
async def create_project(project: Project):
    project_id = str(uuid.uuid4())
    projects[project_id] = {
        "id": project_id,
        "name": project.name,
        "description": project.description,
        "created_at": datetime.now().isoformat()
    }
    save_projects()
    return {"project_id": project_id, "project": projects[project_id]}

@app.get("/list_projects/")
async def list_projects():
    return list(projects.values())

@app.delete("/delete_project/{project_id}")
async def delete_project(project_id: str):
    if project_id in projects:
        # Move all chats from this project to no project
        for chat_id, chat in chat_metadata.items():
            if chat.get("project_id") == project_id:
                chat["project_id"] = None
        del projects[project_id]
        save_projects()
        save_chats()
        return {"message": "Project deleted"}
    raise HTTPException(status_code=404, detail="Project not found")

# Enhanced chat endpoints
@app.post("/new_chat/")
async def new_chat(project_id: Optional[str] = None):
    session_id = str(uuid.uuid4())
    chat_sessions[session_id] = {"summary": "", "messages": []}
    
    chat_metadata[session_id] = {
        "session_id": session_id,
        "name": "New Chat",
        "project_id": project_id,
        "is_favorite": False,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat()
    }
    save_chats()
    return {"session_id": session_id, "metadata": chat_metadata[session_id]}

@app.post("/rename_chat/")
async def rename_chat(request: RenameRequest):
    if request.session_id in chat_metadata:
        chat_metadata[request.session_id]["name"] = request.new_name
        chat_metadata[request.session_id]["updated_at"] = datetime.now().isoformat()
        save_chats()
        return {"message": "Chat renamed successfully"}
    raise HTTPException(status_code=404, detail="Chat not found")

@app.post("/move_to_project/")
async def move_to_project(request: MoveToProjectRequest):
    if request.session_id in chat_metadata:
        chat_metadata[request.session_id]["project_id"] = request.project_id
        chat_metadata[request.session_id]["updated_at"] = datetime.now().isoformat()
        save_chats()
        return {"message": "Chat moved successfully"}
    raise HTTPException(status_code=404, detail="Chat not found")

@app.post("/toggle_favorite/")
async def toggle_favorite(request: ToggleFavoriteRequest):
    if request.session_id in chat_metadata:
        chat_metadata[request.session_id]["is_favorite"] = not chat_metadata[request.session_id]["is_favorite"]
        chat_metadata[request.session_id]["updated_at"] = datetime.now().isoformat()
        save_chats()
        return {"is_favorite": chat_metadata[request.session_id]["is_favorite"]}
    raise HTTPException(status_code=404, detail="Chat not found")

@app.get("/list_chats/")
async def list_chats():
    """Return organized chat list with projects."""
    organized_chats = {
        "favorites": [],
        "projects": {},
        "no_project": []
    }
    
    # Sort projects by creation date (newest first)
    sorted_projects = sorted(projects.items(), key=lambda x: x[1]['created_at'], reverse=True)
    
    # Initialize projects in sorted order
    for project_id, project_data in sorted_projects:
        organized_chats["projects"][project_id] = {
            "project": project_data,
            "chats": []
        }
    
    # Get chat summaries
    for session_id, metadata in chat_metadata.items():
        chat_info = {
            **metadata,
            "summary": chat_sessions.get(session_id, {}).get("summary", "New Chat")
        }
        
        if metadata["is_favorite"]:
            organized_chats["favorites"].append(chat_info)
        
        if metadata["project_id"] and metadata["project_id"] in organized_chats["projects"]:
            organized_chats["projects"][metadata["project_id"]]["chats"].append(chat_info)
        elif not metadata["project_id"]:
            organized_chats["no_project"].append(chat_info)
    
    # Sort chats within each category by updated_at (newest first)
    organized_chats["favorites"].sort(key=lambda x: x["updated_at"], reverse=True)
    organized_chats["no_project"].sort(key=lambda x: x["updated_at"], reverse=True)
    
    for project_id in organized_chats["projects"]:
        organized_chats["projects"][project_id]["chats"].sort(
            key=lambda x: x["updated_at"], reverse=True
        )
    
    return organized_chats

# Enhanced generate endpoint with citations
@app.post("/generate/")
async def generate_response(request: QueryRequest):
    session_id = request.session_id

    if session_id not in chat_sessions:
        chat_sessions[session_id] = {"summary": "", "messages": []}

    chat_sessions[session_id]["messages"].append({"role": "user", "message": request.query})

    # Update chat name if it's the first message
    if not chat_sessions[session_id]["summary"]:
        chat_sessions[session_id]["summary"] = request.query[:30] + ("..." if len(request.query) > 30 else "")
        if session_id in chat_metadata and chat_metadata[session_id]["name"] == "New Chat":
            chat_metadata[session_id]["name"] = chat_sessions[session_id]["summary"]
            save_chats()

    # Generate response with citations
    response_data = rag_model.generate_response_with_citations(
        request.query, 
        history=chat_sessions[session_id]["messages"]
    )
    
    chat_sessions[session_id]["messages"].append({
        "role": "assistant", 
        "message": response_data['response'],
        "citations": response_data.get('citations', {}),
        "highlighted_passages": response_data.get('highlighted_passages', {})
    })
    
    # Update last updated time
    if session_id in chat_metadata:
        chat_metadata[session_id]["updated_at"] = datetime.now().isoformat()
        save_chats()
    
    return {
        "session_id": session_id, 
        "answer": response_data['response'],
        "citations": response_data.get('citations', {}),
        "highlighted_passages": response_data.get('highlighted_passages', {})
    }

# Document viewing endpoint
@app.post("/view_document/")
async def view_document(request: DocumentViewRequest):
    """Return document content with highlight positions."""
    try:
        # Read the text file
        text_path = Path(config.TXT_DIRECTORY) / "Text" / request.filename
        if not text_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")
        
        with open(text_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        # Process highlights with improved matching
        highlights = []
        for highlight in request.highlights:
            passage = highlight.get('passage', '').strip()
            if not passage:
                continue
                
            # Try exact match first
            start = content.find(passage)
            if start != -1:
                highlights.append({
                    'start': start,
                    'end': start + len(passage),
                    'passage': passage
                })
            else:
                # Try normalized match (remove extra spaces, newlines)
                normalized_content = ' '.join(content.split())
                normalized_passage = ' '.join(passage.split())
                norm_start = normalized_content.find(normalized_passage)
                
                if norm_start != -1:
                    # Map back to original position
                    char_count = 0
                    orig_pos = 0
                    for i, char in enumerate(content):
                        if not char.isspace() or (i > 0 and not content[i-1].isspace()):
                            if char_count == norm_start:
                                orig_pos = i
                                break
                            char_count += 1
                    
                    # Find approximate end position
                    end_pos = orig_pos + len(passage)
                    highlights.append({
                        'start': orig_pos,
                        'end': end_pos,
                        'passage': passage,
                        'approximate': True
                    })
        
        return {
            'content': content,
            'highlights': highlights,
            'filename': request.filename
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Enhanced export functionality
def clean_citations_from_text(text: str) -> tuple[str, List[str]]:
    """Remove inline citations and extract unique sources."""
    # Remove citations like [1], [2], etc.
    clean_text = re.sub(r'\[\d+\]', '', text)
    
    # Extract citation numbers
    citation_numbers = re.findall(r'\[(\d+)\]', text)
    
    return clean_text, list(set(citation_numbers))

def create_docx_from_message(message: Dict) -> bytes:
    """Create a DOCX file from a single message with proper formatting."""
    doc = DocxDocument()
    
    # Add title
    title = doc.add_heading('AI Response', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")
    
    # Get clean content and citation numbers
    clean_content, citation_nums = clean_citations_from_text(message['message'])
    
    # Process the content by paragraphs
    paragraphs = clean_content.split('\n\n')
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        # Check if it's a header (simple heuristic)
        if para.strip().endswith(':') and len(para.strip()) < 50:
            p = doc.add_heading(para.strip(), level=2)
        elif para.strip().startswith('- ') or para.strip().startswith('• '):
            # Handle bullet points
            p = doc.add_paragraph(para.strip(), style='List Bullet')
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            
            # Process bold text
            parts = re.split(r'(\*\*.*?\*\*)', para)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Normal text
                    p.add_run(part)
    
    # Add sources section if there are citations
    if citation_nums and 'citations' in message:
        doc.add_page_break()
        doc.add_heading('Sources', 1)
        
        unique_sources = []
        for num in sorted(citation_nums, key=int):
            if num in message['citations']:
                citation = message['citations'][num]
                source_text = f"[{num}] {citation['filename']}"
                if citation['source_url']:
                    source_text += f" - {citation['source_url']}"
                if source_text not in unique_sources:
                    unique_sources.append(source_text)
                    doc.add_paragraph(source_text)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_pdf_from_message(message: Dict) -> bytes:
    """Create a PDF file from a single message with proper formatting."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading2']
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY
    )
    
    story = []
    
    # Add title
    story.append(Paragraph("AI Response", title_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Add metadata
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Get clean content and citation numbers
    clean_content, citation_nums = clean_citations_from_text(message['message'])
    
    # Process content
    paragraphs = clean_content.split('\n\n')
    
    for para in paragraphs:
        if not para.strip():
            continue
            
        # Clean up the paragraph for PDF
        para = para.replace('**', '<b>').replace('**', '</b>')
        para = para.replace('\n', '<br/>')
        
        if para.strip().endswith(':') and len(para.strip()) < 50:
            story.append(Paragraph(para.strip(), heading_style))
        else:
            story.append(Paragraph(para, normal_style))
        story.append(Spacer(1, 0.2*inch))
    
    # Add sources if there are citations
    if citation_nums and 'citations' in message:
        story.append(PageBreak())
        story.append(Paragraph("Sources", heading_style))
        story.append(Spacer(1, 0.2*inch))
        
        unique_sources = []
        for num in sorted(citation_nums, key=int):
            if num in message['citations']:
                citation = message['citations'][num]
                source_text = f"[{num}] {citation['filename']}"
                if citation['source_url']:
                    source_text += f" - {citation['source_url']}"
                if source_text not in unique_sources:
                    unique_sources.append(source_text)
                    story.append(Paragraph(source_text, normal_style))
                    story.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

@app.post("/export_chat/")
async def export_chat(request: ExportRequest):
    """Export a specific message in the requested format."""
    session_id = request.session_id
    
    if session_id not in chat_sessions:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    messages = chat_sessions[session_id]["messages"]
    
    if request.message_index >= len(messages):
        raise HTTPException(status_code=404, detail="Message not found")
    
    message = messages[request.message_index]
    
    if message['role'] != 'assistant':
        raise HTTPException(status_code=400, detail="Can only export AI responses")
    
    if request.format == "docx":
        content = create_docx_from_message(message)
        filename = f"ai_response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif request.format == "pdf":
        content = create_pdf_from_message(message)
        filename = f"ai_response_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        media_type = "application/pdf"
    else:
        raise HTTPException(status_code=400, detail="Invalid export format")
    
    return Response(
        content=content,
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

# Original endpoints maintained for compatibility
@app.post("/get_chat_history/")
async def get_chat_history(request: ChatHistoryRequest):
    session_id = request.session_id
    history = chat_sessions.get(session_id, {"messages": []})["messages"]
    return {"session_id": session_id, "history": history}

@app.post("/save_rating/")
async def save_rating(request: RatingRequest):
    try:
        with open(RATINGS_FILE, "r") as f:
            ratings = json.load(f)
        
        ratings.append({
            "question": request.question,
            "response": request.response,
            "feedback_type": request.feedback_type,
            "selected_reason": request.selected_reason,
            "custom_feedback": request.custom_feedback,
            "timestamp": datetime.now().isoformat()
        })
        
        with open(RATINGS_FILE, "w") as f:
            json.dump(ratings, f, indent=4)
        
        return {"message": "Rating saved successfully"}
    except Exception as e:
        return {"error": str(e)}

# Helper functions
def save_projects():
    with open(PROJECTS_FILE, "w") as f:
        json.dump(projects, f, indent=4)

def save_chats():
    with open(CHATS_FILE, "w") as f:
        json.dump(chat_metadata, f, indent=4)

# Cleanup endpoint for temporary files
@app.on_event("startup")
async def startup_event():
    # Clean up old export files
    for file in DATA_DIR.glob("export_*.*"):
        if file.is_file():
            file.unlink()

@app.delete("/delete_chat/{session_id}")
async def delete_chat(session_id: str):
    if session_id in chat_sessions:
        del chat_sessions[session_id]
    if session_id in chat_metadata:
        del chat_metadata[session_id]
    save_chats()
    return {"message": "Chat deleted successfully"}
